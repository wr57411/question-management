from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_test_documents():
    """创建测试用的Word文档"""
    # 创建题目文档
    questions_doc = Document()
    
    # 添加标题
    title = questions_doc.add_paragraph("物理测试题")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True
    
    # 添加选择题
    questions_doc.add_paragraph("1")
    questions_doc.add_paragraph("一个物体做匀速直线运动，下列说法正确的是")
    questions_doc.add_paragraph("A. 物体的速度保持不变\nB. 物体的加速度为零\nC. 物体受到的合外力为零\nD. 以上说法都正确")
    
    # 添加计算题
    questions_doc.add_paragraph("2")
    questions_doc.add_paragraph("计算题：一个质量为2kg的物体，受到10N的力，求其加速度。")
    
    # 添加填空题
    questions_doc.add_paragraph("3")
    questions_doc.add_paragraph("填空题：牛顿第一定律指出，物体不受力时，将保持（    ）状态或（    ）状态。")
    
    # 保存题目文档
    questions_doc.save("test_questions.docx")
    print("已创建题目文档：test_questions.docx")
    
    # 创建答案文档
    answers_doc = Document()
    
    # 添加标题
    title = answers_doc.add_paragraph("物理测试题答案")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(16)
    title.runs[0].font.bold = True
    
    # 添加选择题答案
    answers_doc.add_paragraph("1")
    answers_doc.add_paragraph("答案：D")
    answers_doc.add_paragraph("解析：匀速直线运动的特点是速度大小和方向都不变，加速度为零，根据牛顿第一定律，物体受到的合外力为零。")
    
    # 添加计算题答案
    answers_doc.add_paragraph("2")
    answers_doc.add_paragraph("答案：5m/s²")
    answers_doc.add_paragraph("解析：根据牛顿第二定律F=ma，a=F/m=10N/2kg=5m/s²")
    
    # 添加填空题答案
    answers_doc.add_paragraph("3")
    answers_doc.add_paragraph("答案：静止；匀速直线运动")
    answers_doc.add_paragraph("解析：牛顿第一定律的基本内容")
    
    # 保存答案文档
    answers_doc.save("test_answers.docx")
    print("已创建答案文档：test_answers.docx")

if __name__ == "__main__":
    create_test_documents() 
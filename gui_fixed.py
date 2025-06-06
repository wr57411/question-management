import sys
import os
import requests
import json
import re
import traceback
from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QPushButton, QFileDialog, QTextEdit, QVBoxLayout, 
    QHBoxLayout, QWidget, QLabel, QTableWidget, QTableWidgetItem, QMessageBox, 
    QHeaderView, QSplitter, QGroupBox, QComboBox, QScrollArea
)
from PyQt5.QtCore import Qt, QBuffer
from PyQt5.QtGui import QTextDocument, QTextCursor, QPixmap, QImage
import docx
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
from PIL import Image
import base64
import os
try:
    from PyQt5.QtWebEngineWidgets import QWebEngineView
    WEB_ENGINE_AVAILABLE = True
    print("PyQtWebEngine 已成功加�?)
except ImportError:
    WEB_ENGINE_AVAILABLE = False
    print("警告：PyQtWebEngine 未安装，图片预览可能无法正常显示")

print("正在初始化应用程�?..")

# 自定义Python执行器包装类
class CustomPythonExecutor:
    @staticmethod
    def run(func, *args, **kwargs):
        try:
            return func(*args, **kwargs)
        except Exception as e:
            err_msg = str(e)
            if "did not find executable at 'C:\\Python313\\python.exe'" in err_msg:
                print("检测到Python路径错误，尝试修�?..")
                # 自动修复Python路径问题
                CustomPythonExecutor.fix_python_path()
                try:
                    # 再次尝试运行
                    return func(*args, **kwargs)
                except Exception as e2:
                    print(f"修复后再次尝试运行时出错: {str(e2)}")
                    traceback.print_exc()
                    return None
            else:
                print(f"执行时出�? {err_msg}")
                traceback.print_exc()
                return None
    
    @staticmethod
    def fix_python_path():
        """尝试修复Python路径问题"""
        try:
            # 检查系统和用户环境变量
            path = os.environ.get("PATH", "")
            
            # 输出当前路径信息，方便调�?            print(f"当前PATH环境变量: {path}")
            
            if "Python313" in path:
                print("检测到错误的Python路径(Python313)，尝试修�?..")
                path_parts = path.split(";")
                new_path_parts = [p for p in path_parts if "Python313" not in p]
                os.environ["PATH"] = ";".join(new_path_parts)
                print("已从PATH中移除错误的Python路径")
                
            # 检查用户变量中的Python310是否存在
            python310_parts = [p for p in path.split(";") if "Python310" in p]
            for part in python310_parts:
                if not os.path.exists(os.path.join(part, "python.exe")):
                    print(f"发现无效的Python310路径: {part}")
                    # 这里我们不移除Python310，因为用户说他有Python310
                    # 但我们会输出警告信息
                    print("警告: 您的环境变量指向不存在的Python路径，请检查环境变量设�?)
                
            # 检查当前目录下是否有可用的Python
            venv_paths = [
                ".venv/Scripts",
                ".venv\\Scripts",
                "venv/Scripts", 
                "venv\\Scripts"
            ]
            
            for venv_path in venv_paths:
                if os.path.exists(f"{venv_path}/python.exe") or os.path.exists(f"{venv_path}\\python.exe"):
                    full_path = os.path.abspath(venv_path)
                    print(f"检测到本地虚拟环境，将其添加到PATH: {full_path}")
                    # 将虚拟环境添加到PATH的最前面
                    os.environ["PATH"] = full_path + ";" + os.environ.get("PATH", "")
                    os.environ["VIRTUAL_ENV"] = os.path.dirname(full_path)
                    print("已添加虚拟环境到PATH")
                    break
                
            return True
        except Exception as e:
            print(f"修复Python路径时出�? {str(e)}")
            traceback.print_exc()
            return False

class QuestionRuleEngine:
    def __init__(self, rules_file='question_rules.json'):
        self.rules = self.load_rules(rules_file)
        self.patterns = self.compile_patterns()
    
    def load_rules(self, rules_file):
        """加载规则配置文件"""
        try:
            with open(rules_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        except Exception as e:
            print(f"加载规则文件失败: {str(e)}")
            return {"rules": [], "content_rules": []}
    
    def compile_patterns(self):
        """编译正则表达式模�?""
        patterns = {
            'question': [],
            'content': []
        }
        
        # 编译题目识别规则
        for rule in self.rules.get('rules', []):
            try:
                pattern = re.compile(rule['pattern'])
                patterns['question'].append({
                    'pattern': pattern,
                    'priority': rule.get('priority', 0),
                    'name': rule.get('name', '')
                })
            except Exception as e:
                print(f"编译规则失败 {rule.get('name', '')}: {str(e)}")
        
        # 编译内容过滤规则
        for rule in self.rules.get('content_rules', []):
            try:
                pattern = re.compile(rule['pattern'])
                patterns['content'].append({
                    'pattern': pattern,
                    'action': rule.get('action', 'exclude'),
                    'name': rule.get('name', '')
                })
            except Exception as e:
                print(f"编译规则失败 {rule.get('name', '')}: {str(e)}")
        
        # 按优先级排序题目识别规则
        patterns['question'].sort(key=lambda x: x['priority'])
        return patterns
    
    def is_new_question(self, text):
        """判断是否是新题目的开�?""
        for rule in self.patterns['question']:
            if rule['pattern'].match(text):
                return True
        return False
    
    def should_exclude(self, text):
        """判断是否应该排除该内�?""
        for rule in self.patterns['content']:
            if rule['pattern'].match(text):
                return rule['action'] == 'exclude'
        return False

class PhysicsQuestionGUI(QMainWindow):
    def __init__(self):
        super().__init__()
        self.setWindowTitle("物理题目管理系统")
        self.setGeometry(100, 100, 1200, 800)
        
        # 初始化规则引�?        self.rule_engine = QuestionRuleEngine()
        
        # 创建主布局
        main_layout = QVBoxLayout()
        
        # 创建水平分割�?        splitter = QSplitter(Qt.Horizontal)
        
        # 左侧面板 - 文档预览
        left_panel = QWidget()
        left_layout = QVBoxLayout()
        
        # 文档上传部分
        upload_group = QGroupBox("文档上传")
        upload_layout = QVBoxLayout()
        
        self.upload_btn = QPushButton("上传题目文档", self)
        self.upload_btn.clicked.connect(self.upload_questions)
        upload_layout.addWidget(self.upload_btn)
        
        # 文档预览（使用QScrollArea包装�?        scroll_area = QScrollArea()
        scroll_area.setWidgetResizable(True)
        self.doc_preview = QTextEdit(self)
        self.doc_preview.setReadOnly(True)
        scroll_area.setWidget(self.doc_preview)
        upload_layout.addWidget(QLabel("文档预览�?))
        upload_layout.addWidget(scroll_area)
        
        upload_group.setLayout(upload_layout)
        left_layout.addWidget(upload_group)
        
        # 题目列表
        table_group = QGroupBox("题目列表")
        table_layout = QVBoxLayout()
        
        self.table = QTableWidget(self)
        self.table.setColumnCount(6)  # 增加一列用于粘贴答案按�?        self.table.setHorizontalHeaderLabels(["题号", "题目", "答案", "解析", "粘贴答案", "操作"])
        self.table.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        # 设置表格样式
        self.table.setStyleSheet("""
            QTableWidget {
                gridline-color: #d0d0d0;
                background-color: white;
                alternate-background-color: #f6f6f6;
            }
            QTableWidget::item {
                padding: 5px;
            }
            QPushButton {
                padding: 5px;
                margin: 2px;
            }
        """)
        table_layout.addWidget(self.table)
        
        table_group.setLayout(table_layout)
        left_layout.addWidget(table_group)
        
        left_panel.setLayout(left_layout)
        
        # 右侧面板 - 答案编辑
        right_panel = QWidget()
        right_layout = QVBoxLayout()
        
        # 题目预览部分
        preview_group = QGroupBox("题目预览")
        preview_layout = QVBoxLayout()
        
        if WEB_ENGINE_AVAILABLE:
            self.question_preview = QWebEngineView(self)
        else:
            self.question_preview = QTextEdit(self)
            self.question_preview.setReadOnly(True)
            self.question_preview.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOn)
            self.question_preview.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOn)
            self.question_preview.setStyleSheet("""
                QTextEdit {
                    border: 1px solid #ccc;
                    border-radius: 4px;
                    padding: 5px;
                }
                QScrollBar:vertical {
                    border: none;
                    background: #f0f0f0;
                    width: 10px;
                    margin: 0px;
                }
                QScrollBar::handle:vertical {
                    background: #c0c0c0;
                    min-height: 20px;
                    border-radius: 5px;
                }
                QScrollBar::add-line:vertical, QScrollBar::sub-line:vertical {
                    height: 0px;
                }
                QScrollBar:horizontal {
                    border: none;
                    background: #f0f0f0;
                    height: 10px;
                    margin: 0px;
                }
                QScrollBar::handle:horizontal {
                    background: #c0c0c0;
                    min-width: 20px;
                    border-radius: 5px;
                }
                QScrollBar::add-line:horizontal, QScrollBar::sub-line:horizontal {
                    width: 0px;
                }
            """)
        preview_layout.addWidget(self.question_preview)
        
        preview_group.setLayout(preview_layout)
        right_layout.addWidget(preview_group)
        
        # 答案编辑部分
        answer_group = QGroupBox("答案编辑")
        answer_layout = QVBoxLayout()
        
        # 答案类型选择
        type_layout = QHBoxLayout()
        type_layout.addWidget(QLabel("答案类型�?))
        self.answer_type = QComboBox()
        self.answer_type.addItems(["选择�?, "计算�?, "填空�?, "问答�?])
        type_layout.addWidget(self.answer_type)
        answer_layout.addLayout(type_layout)
        
        # 答案输入
        self.answer_text = QTextEdit(self)
        self.answer_text.setPlaceholderText("请粘贴答案文本，支持Word格式")
        answer_layout.addWidget(QLabel("答案输入�?))
        answer_layout.addWidget(self.answer_text)
        
        # 答案预览
        self.answer_preview = QTextEdit(self)
        self.answer_preview.setReadOnly(True)
        answer_layout.addWidget(QLabel("答案预览�?))
        answer_layout.addWidget(self.answer_preview)
        
        # 提交按钮
        self.submit_btn = QPushButton("提交", self)
        self.submit_btn.clicked.connect(self.submit)
        answer_layout.addWidget(self.submit_btn)
        
        answer_group.setLayout(answer_layout)
        right_layout.addWidget(answer_group)
        
        right_panel.setLayout(right_layout)
        
        # 添加面板到分割器
        splitter.addWidget(left_panel)
        splitter.addWidget(right_panel)
        splitter.setSizes([600, 400])  # 设置初始大小
        
        main_layout.addWidget(splitter)
        
        # 设置主布局
        central_widget = QWidget()
        central_widget.setLayout(main_layout)
        self.setCentralWidget(central_widget)
        
        self.questions_file = None
        # 连接表格选择变化信号
        self.table.itemSelectionChanged.connect(self.update_question_preview)
        self.refresh_questions()
    
    def upload_questions(self):
        file_name, _ = QFileDialog.getOpenFileName(self, "选择题目文档", "", "Word Files (*.docx)")
        if file_name:
            self.questions_file = file_name
            try:
                # 读取Word文档
                doc = docx.Document(file_name)
                
                # 显示完整文档内容在预览框�?                preview_text = "<html><body>"
                for para in doc.paragraphs:
                    if para.style.name.startswith('Header') or para.style.name.startswith('Footer'):
                        continue
                    if para.text.strip():
                        preview_text += f"<p>{para.text}</p>"
                    
                    # 处理段落中的图片
                    for run in para.runs:
                        if run._element.xpath('.//w:drawing'):
                            for shape in run._element.xpath('.//w:drawing//a:blip'):
                                embed_id = shape.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if embed_id:
                                    try:
                                        image_part = doc.part.related_parts[embed_id]
                                        image_data = image_part.blob
                                        image_base64 = base64.b64encode(image_data).decode()
                                        preview_text += f'<img src="data:image/png;base64,{image_base64}" style="max-width: 100%; height: auto;" />'
                                    except Exception as e:
                                        print(f"处理图片时出�? {str(e)}")
                
                preview_text += "</body></html>"
                self.doc_preview.setHtml(preview_text)
                
                # 清空表格
                self.table.setRowCount(0)
                
                # 用于存储当前题目内容和图�?                current_question = ""
                current_images = []
                question_number = 1
                processed_images = set()  # 用于跟踪已处理的图片
                
                print("\n开始处理文档内容：")
                print("-" * 50)
                
                # 获取所有段�?                paragraphs = [p for p in doc.paragraphs if not p.style.name.startswith(('Header', 'Footer'))]
                total_paragraphs = len(paragraphs)
                
                # 遍历文档段落
                for current_para_index, para in enumerate(paragraphs):
                    text = para.text.strip()
                    
                    # 检查是否应该排除该内容
                    if self.rule_engine.should_exclude(text):
                        print(f"排除内容: {text}")
                        continue
                    
                    print(f"\n处理段落 {current_para_index + 1}/{total_paragraphs}: {text}")
                    
                    # 检查是否是新题�?                    is_new = self.rule_engine.is_new_question(text)
                    print(f"是否为新题目: {is_new}")
                    
                    if is_new:
                        # 添加之前的题目到表格
                        if current_question:
                            print(f"\n添加题目 {question_number} 到表�?")
                            print(f"题目内容: {current_question}")
                            print(f"当前题目包含 {len(current_images)} 张图�?)
                            # 将图片添加到题目内容�?                            question_with_images = current_question
                            for img in current_images:
                                question_with_images += f"\n{img}"
                            print("添加题目到表格，包含图片")
                            self.add_question_to_table(question_number, question_with_images)
                            question_number += 1
                        current_question = text
                        current_images = []  # 清空图片列表
                        processed_images.clear()  # 清空已处理图片集�?                        print(f"开始新题目: {text}")
                    else:
                        # 继续累积当前题目内容
                        if current_question:
                            current_question += "\n" + text
                            print(f"累积题目内容: {text}")
                        else:
                            current_question = text
                            print(f"开始第一个题�? {text}")
                    
                    # 处理段落中的图片
                    for run in para.runs:
                        if run._element.xpath('.//w:drawing'):
                            print("在当前段落中发现图片")
                            # 获取图片数据
                            for shape in run._element.xpath('.//w:drawing//a:blip'):
                                embed_id = shape.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if embed_id and embed_id not in processed_images:
                                    try:
                                        print(f"处理图片ID: {embed_id}")
                                        processed_images.add(embed_id)  # 标记图片已处�?                                        image_part = doc.part.related_parts[embed_id]
                                        image_data = image_part.blob
                                        
                                        # 将图片数据转换为base64
                                        image_base64 = base64.b64encode(image_data).decode()
                                        # 创建图片HTML标签
                                        img_html = f'<img src="data:image/png;base64,{image_base64}" style="max-width: 100%; height: auto;" />'
                                        # 将图片添加到当前题目的图片列表中
                                        current_images.append(img_html)
                                        print(f"成功添加图片到当前题目，当前题目共有 {len(current_images)} 张图�?)
                                    except Exception as e:
                                        print(f"处理图片时出�? {str(e)}")
                    
                    # 处理段落后的图片（作为独立对象）
                    # 只在当前段落是新题目时处理图片，并且只处理下一张图�?                    if is_new and current_para_index + 1 < total_paragraphs:
                        next_para = paragraphs[current_para_index + 1]
                        next_text = next_para.text.strip()
                        
                        # 如果下一个段落存在且不是新题目，则处理其中的图片
                        if next_text and not self.rule_engine.is_new_question(next_text):
                            for run in next_para.runs:
                                if run._element.xpath('.//w:drawing'):
                                    print("在下一段落中发现图�?)
                                    for shape in run._element.xpath('.//w:drawing//a:blip'):
                                        embed_id = shape.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                        if embed_id and embed_id not in processed_images:
                                            try:
                                                print(f"处理图片ID: {embed_id}")
                                                processed_images.add(embed_id)
                                                image_part = doc.part.related_parts[embed_id]
                                                image_data = image_part.blob
                                                image_base64 = base64.b64encode(image_data).decode()
                                                img_html = f'<img src="data:image/png;base64,{image_base64}" style="max-width: 100%; height: auto;" />'
                                                current_images.append(img_html)
                                                print(f"成功添加图片到当前题目，当前题目共有 {len(current_images)} 张图�?)
                                                break  # 只添加一张图�?                                            except Exception as e:
                                                print(f"处理图片时出�? {str(e)}")
                
                # 添加最后一个题�?                if current_question:
                    print(f"\n添加最后一个题�?{question_number} 到表�?")
                    print(f"题目内容: {current_question}")
                    print(f"最后一个题目包�?{len(current_images)} 张图�?)
                    # 将图片添加到题目内容�?                    question_with_images = current_question
                    for img in current_images:
                        question_with_images += f"\n{img}"
                    print("添加最后一个题目到表格，包含图�?)
                    self.add_question_to_table(question_number, question_with_images)
                
                print("\n文档处理完成")
                print("-" * 50)
                
                QMessageBox.information(self, "提示", f"已加载题目文档：{file_name}")
                
            except Exception as e:
                QMessageBox.critical(self, "错误", f"读取文档失败：{str(e)}")
                print(f"错误详情: {str(e)}")
    
    def add_question_to_table(self, question_number, content):
        """添加题目到表格中"""
        row = self.table.rowCount()
        self.table.insertRow(row)
        
        # 设置题号
        self.table.setItem(row, 0, QTableWidgetItem(str(question_number)))
        
        # 设置题目内容，保留HTML格式
        content_item = QTableWidgetItem()
        content_item.setData(Qt.UserRole, content)  # 用UserRole存储HTML
        self.table.setItem(row, 1, content_item)
        
        # 设置空的答案和解�?        self.table.setItem(row, 2, QTableWidgetItem(""))
        self.table.setItem(row, 3, QTableWidgetItem(""))
        
        # 添加粘贴答案按钮
        paste_btn = QPushButton("粘贴答案", self)
        paste_btn.setStyleSheet("""
            QPushButton {
                background-color: #4CAF50;
                color: white;
                border: none;
                padding: 5px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #45a049;
            }
        """)
        paste_btn.clicked.connect(lambda checked, q_id=question_number: self.paste_answer_for_question(q_id))
        self.table.setCellWidget(row, 4, paste_btn)
        
        # 添加删除按钮
        delete_btn = QPushButton("删除", self)
        delete_btn.setStyleSheet("""
            QPushButton {
                background-color: #f44336;
                color: white;
                border: none;
                padding: 5px;
                border-radius: 3px;
            }
            QPushButton:hover {
                background-color: #da190b;
            }
        """)
        delete_btn.clicked.connect(lambda checked, q_id=question_number: self.delete_question(q_id))
        self.table.setCellWidget(row, 5, delete_btn)
        
        # 调整列宽
        self.table.resizeColumnsToContents()
        self.table.setColumnWidth(4, 100)  # 粘贴答案按钮�?        self.table.setColumnWidth(5, 80)   # 删除按钮�?    
    def update_question_preview(self):
        """更新题目预览内容"""
        selected_items = self.table.selectedItems()
        if not selected_items:
            return
            
        # 获取选中行的索引
        row = selected_items[0].row()
        
        # 获取题目内容
        question_content = self.table.item(row, 1).data(Qt.UserRole)
        question_number = self.table.item(row, 0).text()
        answer = self.table.item(row, 2).text()
        analysis = self.table.item(row, 3).text()
        
        print(f"\n更新题目预览:")
        print(f"题目编号: {question_number}")
        print(f"题目内容长度: {len(question_content)}")
        print(f"是否包含图片标签: {'<img' in question_content}")
        
        # 测试用base64图片�?x1红色像素PNG�?        test_img = '<img src="data:image/png;base64,iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/w8AAn8B9pQn2wAAAABJRU5ErkJggg==" style="max-width:100%;height:auto;" />'
        preview_html = f"""
        <html>
        <head>
            <style>
                img {{ 
                    max-width: 100%; 
                    height: auto; 
                    display: block;
                    margin: 10px 0;
                }}
                .content-block {{ 
                    margin: 10px; 
                    padding: 10px; 
                    border: 1px solid #ddd; 
                    border-radius: 5px; 
                }}
                .content-block p {{ 
                    margin: 5px 0; 
                }}
                .content-block div {{ 
                    margin-left: 20px; 
                }}
            </style>
        </head>
        <body>
            {test_img if '<img' in (question_content or '') else ''}
            <h3>第{question_number}�?/h3>
            <div class="content-block">
                <p><b>题目内容�?/b></p>
                <div>{question_content}</div>
            </div>
        """
        
        if answer:
            preview_html += f"""
            <div class="content-block">
                <p><b>答案�?/b></p>
                <div>{answer}</div>
            </div>
            """
            
        if analysis:
            preview_html += f"""
            <div class="content-block">
                <p><b>解析�?/b></p>
                <div>{analysis}</div>
            </div>
            """
            
        preview_html += """
        </body>
        </html>
        """
        
        print("设置HTML预览内容")
        # 设置HTML内容
        if WEB_ENGINE_AVAILABLE:
            self.question_preview.setHtml(preview_html)
        else:
            self.question_preview.setHtml(preview_html)
            if not WEB_ENGINE_AVAILABLE:
                print("警告：未安装 PyQtWebEngine，图片预览可能无法正常显示。请安装 PyQtWebEngine 获得最佳体验�?)
    
    def paste_answer_for_question(self, question_id):
        try:
            # 获取剪贴板内�?            clipboard = QApplication.clipboard()
            mime_data = clipboard.mimeData()
            
            # 创建临时Word文档
            doc = docx.Document()
            
            # 处理文本内容
            if mime_data.hasText():
                text = mime_data.text()
                doc.add_paragraph(text)
            
            # 处理图片内容
            if mime_data.hasImage():
                print("检测到剪贴板中的图�?)
                image = QImage(mime_data.imageData())
                if not image.isNull():
                    # 将QImage转换为PIL Image
                    buffer = QBuffer()
                    buffer.open(QBuffer.ReadWrite)
                    image.save(buffer, "PNG")
                    pil_image = Image.open(io.BytesIO(buffer.data()))
                    
                    # 保存为临时文�?                    temp_image_path = "temp_image.png"
                    pil_image.save(temp_image_path)
                    
                    # 将图片添加到Word文档
                    doc.add_picture(temp_image_path)
                    
                    # 删除临时文件
                    os.remove(temp_image_path)
            
            # 保存为临时文�?            temp_file = "temp_answer.docx"
            doc.save(temp_file)
            
            # 读取临时文件内容
            doc = docx.Document(temp_file)
            formatted_text = ""
            
            # 处理文档内容
            for para in doc.paragraphs:
                if para.text.strip():
                    formatted_text += para.text + "\n"
                
                # 处理段落中的图片
                for run in para.runs:
                    if run._element.xpath('.//w:drawing'):
                        for shape in run._element.xpath('.//w:drawing//a:blip'):
                            embed_id = shape.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                            if embed_id:
                                image_part = doc.part.related_parts[embed_id]
                                image_data = image_part.blob
                                image_base64 = base64.b64encode(image_data).decode()
                                formatted_text += f'<img src="data:image/png;base64,{image_base64}" style="max-width: 100%; height: auto;" />\n'
            
            # 更新答案输入�?            self.answer_text.setHtml(formatted_text)
            
            # 显示在预览区�?            self.answer_preview.setHtml(formatted_text)
            
            # 删除临时文件
            os.remove(temp_file)
            
        except Exception as e:
            QMessageBox.critical(self, "错误", f"粘贴答案失败：{str(e)}")
            print(f"错误详情: {str(e)}")
    
    def submit(self):
        if not self.questions_file:
            QMessageBox.warning(self, "警告", "请先上传题目文档")
            return
        
        answer_text = self.answer_text.toPlainText()
        if not answer_text.strip():
            QMessageBox.warning(self, "警告", "请输入答案内�?)
            return
        
        try:
            # 准备文件和数�?            files = {
                'file': (os.path.basename(self.questions_file), open(self.questions_file, 'rb'), 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            }
            data = {
                'answer_text': answer_text,
                'gui_mode': 'true'
            }
            
            print("正在提交答案...")
            print(f"文件: {self.questions_file}")
            print(f"答案内容长度: {len(answer_text)}")
            
            # 发送请�?            response = requests.post(
                'http://localhost:8000/api/physics-questions/import-word',
                files=files,
                data=data
            )
            
            # 检查响应状�?            if response.status_code == 200:
                try:
                    result = response.json()
                    if isinstance(result, dict) and 'message' in result:
                        QMessageBox.information(self, "成功", result['message'])
                    else:
                        QMessageBox.information(self, "成功", "答案已成功提�?)
                except json.JSONDecodeError:
                    QMessageBox.information(self, "成功", "答案已成功提�?)
                self.refresh_questions()
            else:
                try:
                    error_detail = response.json().get('detail', '未知错误')
                    QMessageBox.critical(self, "错误", f"上传失败：{error_detail}")
                except json.JSONDecodeError:
                    QMessageBox.critical(self, "错误", f"上传失败：HTTP {response.status_code}")
                
        except requests.exceptions.ConnectionError:
            QMessageBox.critical(self, "错误", "无法连接到服务器，请确保后端服务已启�?)
        except Exception as e:
            QMessageBox.critical(self, "错误", f"上传失败：{str(e)}")
            print(f"错误详情: {str(e)}")
        finally:
            # 确保文件被关�?            if 'files' in locals():
                files['file'][1].close()
    
    def refresh_questions(self):
        try:
            response = requests.get('http://localhost:8000/api/physics-questions/')
            if response.status_code != 200:
                QMessageBox.warning(self, "警告", f"获取题目失败：{response.text}")
                return
                
            questions = response.json()
            if not isinstance(questions, list):
                QMessageBox.warning(self, "警告", "返回数据格式错误")
                return
                
            self.table.setRowCount(len(questions))
            for i, q in enumerate(questions):
                # 安全获取数据，避免键不存在的情况
                question_number = q.get('metadata', {}).get('question_number', '')
                content = q.get('content', '')
                answer = q.get('answer', '')
                analysis = q.get('analysis', '')
                question_id = q.get('id', '')
                
                # 设置表格内容
                self.table.setItem(i, 0, QTableWidgetItem(str(question_number)))
                self.table.setItem(i, 1, QTableWidgetItem(str(content)))
                self.table.setItem(i, 2, QTableWidgetItem(str(answer)))
                self.table.setItem(i, 3, QTableWidgetItem(str(analysis)))
                
                # 粘贴答案按钮
                paste_btn = QPushButton("粘贴答案", self)
                paste_btn.setStyleSheet("""
                    QPushButton {
                        background-color: #4CAF50;
                        color: white;
                        border: none;
                        padding: 5px;
                        border-radius: 3px;
                    }
                    QPushButton:hover {
                        background-color: #45a049;
                    }
                """)
                paste_btn.clicked.connect(lambda checked, q_id=question_id: self.paste_answer_for_question(q_id))
                self.table.setCellWidget(i, 4, paste_btn)
                
                # 删除按钮
                delete_btn = QPushButton("删除", self)
                delete_btn.setStyleSheet("""
                    QPushButton {
                        background-color: #f44336;
                        color: white;
                        border: none;
                        padding: 5px;
                        border-radius: 3px;
                    }
                    QPushButton:hover {
                        background-color: #da190b;
                    }
                """)
                delete_btn.clicked.connect(lambda checked, q_id=question_id: self.delete_question(q_id))
                self.table.setCellWidget(i, 5, delete_btn)
                
            # 调整列宽以适应内容
            self.table.resizeColumnsToContents()
            self.table.setColumnWidth(4, 100)  # 粘贴答案按钮�?            self.table.setColumnWidth(5, 80)   # 删除按钮�?            
        except requests.exceptions.ConnectionError:
            QMessageBox.critical(self, "错误", "无法连接到服务器，请确保后端服务已启�?)
        except Exception as e:
            QMessageBox.critical(self, "错误", f"刷新失败：{str(e)}")
            print(f"详细错误信息：{str(e)}")  # 打印详细错误信息以便调试
    
    def delete_question(self, question_id):
        try:
            response = requests.delete(f'http://localhost:8000/api/physics-questions/{question_id}')
            if response.status_code == 200:
                QMessageBox.information(self, "成功", "题目已删�?)
                self.refresh_questions()
            else:
                QMessageBox.warning(self, "警告", f"删除失败：{response.json()['detail']}")
        except Exception as e:
            QMessageBox.critical(self, "错误", f"删除失败：{str(e)}")

if __name__ == "__main__":
    try:
        print("开始初始化...")
        
        # 尝试修复Python路径问题
        CustomPythonExecutor.fix_python_path()
        
        print("创建QApplication实例...")
        # 使用自定义执行器包装QApplication创建
        app = CustomPythonExecutor.run(QApplication, sys.argv)
        if app is None:
            # 如果创建失败，尝试使用其他方式创�?            print("尝试使用替代方式创建应用...")
            os.environ["QT_QPA_PLATFORM"] = "windows"
            app = QApplication(sys.argv)
            
        print("创建主窗�?..")
        window = PhysicsQuestionGUI()
        print("显示主窗�?..")
        window.show()
        print("进入事件循环...")
        sys.exit(app.exec_())
    except Exception as e:
        print(f"程序启动时发生错误：{str(e)}")
        print("错误详情�?)
        traceback.print_exc() 